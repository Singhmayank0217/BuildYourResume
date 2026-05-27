import io
import os
import uuid
from typing import Generator

import psycopg2
import pytest
import redis
from celery import Celery
from docx import Document
from fastapi import FastAPI
from httpx import ASGITransport, AsyncClient
from reportlab.pdfgen import canvas
from starlette.datastructures import UploadFile

# Ensure app settings can initialize during test collection.
os.environ.setdefault("DATABASE_URL", "postgresql://smart:smart@localhost:5432/smart_resume")
os.environ.setdefault("JWT_SECRET", "test-secret")
os.environ.setdefault("REDIS_URL", "redis://localhost:6379/0")
os.environ.setdefault("CELERY_BROKER_URL", os.environ["REDIS_URL"])


@pytest.fixture(scope="session")
def database_url() -> str:
    return os.environ["DATABASE_URL"]


@pytest.fixture(scope="session")
def redis_url() -> str:
    return os.environ["REDIS_URL"]


@pytest.fixture(scope="session")
def db_connection(database_url: str) -> Generator:
    try:
        conn = psycopg2.connect(database_url)
        conn.autocommit = True
    except Exception as exc:
        pytest.skip(f"PostgreSQL unavailable: {exc}")

    try:
        yield conn
    finally:
        conn.close()


@pytest.fixture(scope="session")
def temporary_schema(db_connection) -> Generator[str, None, None]:
    schema_name = f"test_{uuid.uuid4().hex[:12]}"
    with db_connection.cursor() as cur:
        cur.execute(f'CREATE SCHEMA "{schema_name}"')
    try:
        yield schema_name
    finally:
        with db_connection.cursor() as cur:
            cur.execute(f'DROP SCHEMA IF EXISTS "{schema_name}" CASCADE')


@pytest.fixture(scope="session")
def redis_client(redis_url: str):
    client = redis.Redis.from_url(redis_url)
    try:
        client.ping()
    except Exception as exc:
        pytest.skip(f"Redis unavailable: {exc}")
    return client


@pytest.fixture(scope="session")
def celery_app_fixture(redis_url: str) -> Celery:
    app = Celery("test_app", broker=redis_url, backend=redis_url)
    app.conf.update(task_always_eager=True, task_store_eager_result=True)
    return app


@pytest.fixture(scope="session")
def fastapi_test_app() -> FastAPI:
    from app.routers import auth, upload, analyze, health, ai_verify

    app = FastAPI()
    app.include_router(auth.router, prefix="/api/auth", tags=["Auth"])
    app.include_router(upload.router, prefix="/api/resume", tags=["Upload"])
    app.include_router(analyze.router, prefix="/api/analyze", tags=["ATS"])
    app.include_router(ai_verify.router)
    app.include_router(health.router)
    return app


@pytest.fixture()
async def async_client(fastapi_test_app: FastAPI):
    transport = ASGITransport(app=fastapi_test_app)
    async with AsyncClient(transport=transport, base_url="http://testserver") as client:
        yield client


@pytest.fixture()
def sample_resume_payload() -> dict:
    return {
        "full_name": "Alex Johnson",
        "email": "alex@example.com",
        "phone": "+1 555 123 4567",
        "summary": "Backend engineer",
        "skills": ["python", "fastapi", "docker"],
        "experience": [
            {"raw": "Built APIs using FastAPI from 2022 to 2024", "start_year": 2022, "end_year": 2024}
        ],
        "education": [{"raw": "B.Sc Computer Science"}],
    }


@pytest.fixture()
def sample_docx_upload() -> UploadFile:
    doc = Document()
    doc.add_paragraph("Alex Johnson")
    doc.add_paragraph("Email: alex@example.com")
    doc.add_paragraph("Skills: Python, FastAPI, Docker")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return UploadFile(filename="resume.docx", file=buffer)


@pytest.fixture()
def sample_pdf_upload() -> UploadFile:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(100, 750, "Alex Johnson")
    pdf.drawString(100, 730, "Email: alex@example.com")
    pdf.drawString(100, 710, "Skills: Python FastAPI Docker")
    pdf.save()
    buffer.seek(0)
    return UploadFile(filename="resume.pdf", file=buffer)
